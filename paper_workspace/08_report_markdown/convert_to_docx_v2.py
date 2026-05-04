from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPORT_DIR = Path(__file__).resolve().parent
ROOT = REPORT_DIR.parents[1]
OUT_DIR = ROOT / "paper_workspace" / "09_final_docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_FILE = OUT_DIR / "AgriMLOps_KTI_Final.docx"

SECTION_FILES = [
    REPORT_DIR / "front_matter" / "kata_pengantar.md",
    REPORT_DIR / "front_matter" / "abstrak.md",
    REPORT_DIR / "front_matter" / "daftar_isi.md",
    REPORT_DIR / "front_matter" / "daftar_gambar.md",
    REPORT_DIR / "front_matter" / "daftar_tabel.md",
    REPORT_DIR / "bab1" / "pendahuluan.md",
    REPORT_DIR / "bab2" / "kajian_pustaka.md",
    REPORT_DIR / "bab3" / "rancangan_dan_metode.md",
    REPORT_DIR / "bab4" / "hasil_dan_pembahasan.md",
    REPORT_DIR / "bab5" / "penutup.md",
    REPORT_DIR / "daftar_pustaka.md",
    REPORT_DIR / "lampiran" / "lampiran.md",
]


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def style_paragraph(paragraph, align=None):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_text_with_inline_format(paragraph, text: str, size=12):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_heading(doc, text: str, level: int):
    if text.upper().startswith("BAB ") or text in ["KATA PENGANTAR", "ABSTRAK", "DAFTAR ISI", "DAFTAR GAMBAR", "DAFTAR TABEL", "DAFTAR PUSTAKA", "LAMPIRAN"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text.upper())
        set_run_font(run, size=12, bold=True)
        style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        set_run_font(run, size=12, bold=True)
        style_paragraph(paragraph)


def add_markdown_table(doc, rows: list[str]):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if all(set(c.replace(":", "")) <= {"-"} for c in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return
    table = doc.add_table(rows=len(parsed), cols=max(len(r) for r in parsed))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D9EAF7")
                cell._tc.get_or_add_tcPr().append(shd)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_text_with_inline_format(paragraph, cell_text, size=10)
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=(r_idx == 0))
    doc.add_paragraph()


def add_image(doc, image_rel_path: str, alt_text: str):
    img_path = REPORT_DIR / image_rel_path
    if not img_path.exists():
        p = doc.add_paragraph()
        add_text_with_inline_format(p, f"[Gambar belum ditemukan: {image_rel_path}]", size=10)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(img_path), width=Inches(5.7))
    except Exception:
        run.add_text(f"[Gagal memasukkan gambar: {image_rel_path}]")
    style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)


def add_caption(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clean = text.strip("*").strip()
    run = paragraph.add_run(clean)
    set_run_font(run, size=10, bold=True)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)


def add_source(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, italic=True)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)


def add_markdown_file(doc, path: Path):
    if not path.exists():
        return
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, size=9)
            style_paragraph(p)
            i += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            add_heading(doc, text, level)
            i += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                add_image(doc, match.group(2), match.group(1))
            i += 1
            continue
        if line.startswith("**Gambar ") or line.startswith("**Tabel "):
            add_caption(doc, line)
            i += 1
            continue
        if line.startswith("Sumber:"):
            add_source(doc, line)
            i += 1
            continue
        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_markdown_table(doc, table_rows)
            continue
        if line.startswith("- ") or re.match(r"\d+\. ", line):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.25)
            add_text_with_inline_format(p, line, size=12)
            style_paragraph(p)
            i += 1
            continue
        p = doc.add_paragraph()
        add_text_with_inline_format(p, line, size=12)
        style_paragraph(p)
        i += 1


def add_cover(doc):
    for text, space_after in [
        ("KARYA TULIS ILMIAH", 14),
        ("IMPLEMENTASI SISTEM DIAGNOSA PENYAKIT TANAMAN DENGAN", 6),
        ("MACHINE LEARNING OPERATIONS (MLOPS) DAN ACTIVE LEARNING", 18),
        ("SUBTEMA: SISTEM CERDAS BERBASIS DEEP LEARNING", 36),
        ("Asal Instansi: INNOVEST", 24),
        ("Disusun Oleh:", 8),
        ("1. Nama Ketua - NIM: [NIM]", 4),
        ("2. Nama Anggota 1 - NIM: [NIM]", 4),
        ("3. Nama Anggota 2 - NIM: [NIM]", 36),
        ("INNOVEST", 4),
        ("JAKARTA", 4),
        ("2026", 4),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, bold=True)
        p.paragraph_format.space_after = Pt(space_after)
    doc.add_page_break()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
    add_cover(doc)
    for idx, file in enumerate(SECTION_FILES):
        add_markdown_file(doc, file)
        if idx < len(SECTION_FILES) - 1:
            doc.add_page_break()
    doc.save(OUTPUT_FILE)
    shutil_target = ROOT / "report" / "AgriMLOps_KTI_Final.docx"
    doc.save(shutil_target)
    print(f"Saved {OUTPUT_FILE}")
    print(f"Saved {shutil_target}")


if __name__ == "__main__":
    main()
