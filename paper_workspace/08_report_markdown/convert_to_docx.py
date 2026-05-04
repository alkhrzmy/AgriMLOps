"""
Script to convert markdown report to DOCX format
Requirements: pip install python-docx markdown
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import markdown

# Set up paths
REPORT_DIR = r"m:\Workspaces\innovest\report"
OUTPUT_FILE = os.path.join(REPORT_DIR, "AgriMLOps_KTI.docx")

def set_font(run, font_name="Times New Roman", size=12, bold=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    # Set for East Asian characters
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_spacing(paragraph, line_spacing=1.15):
    """Set paragraph spacing"""
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(6)

def add_heading(doc, text, level=1):
    """Add heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_font(run, bold=True, size=12)
    set_paragraph_spacing(heading)
    return heading

def add_paragraph(doc, text, bold=False, size=12):
    """Add paragraph with proper formatting"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        set_font(run, bold=bold, size=size)
    set_paragraph_spacing(para)
    return para

def main():
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Cover Page
    doc.add_page_break()
    add_heading(doc, "KARYA TULIS ILMIAH", level=1)
    add_paragraph(doc, "IMPLEMENTASI SISTEM DIAGNOSA PENYAKIT TANAMAN DENGAN", bold=True)
    add_paragraph(doc, "MACHINE LEARNING OPERATIONS (MLOPS) DAN ACTIVE LEARNING", bold=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "SUBTEMA: SISTEM CERDAS BERBASIS DEEP LEARNING", bold=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "")
    add_paragraph(doc, "Asal Instansi: INNOVEST", bold=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "Disusun Oleh:", bold=True)
    add_paragraph(doc, "1. Nama Ketua", bold=True)
    add_paragraph(doc, "   NIM: [NIM]", bold=True)
    add_paragraph(doc, "2. Nama Anggota 1", bold=True)
    add_paragraph(doc, "   NIM: [NIM]", bold=True)
    add_paragraph(doc, "3. Nama Anggota 2", bold=True)
    add_paragraph(doc, "   NIM: [NIM]", bold=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "")
    add_paragraph(doc, "INNOVEST", bold=True)
    add_paragraph(doc, "JAKARTA", bold=True)
    add_paragraph(doc, "2026", bold=True)
    
    doc.add_page_break()
    
    # Read and add front matter
    front_matter_dir = os.path.join(REPORT_DIR, "front_matter")
    
    # Kata Pengantar
    kata_pengantar_path = os.path.join(front_matter_dir, "kata_pengantar.md")
    if os.path.exists(kata_pengantar_path):
        with open(kata_pengantar_path, 'r', encoding='utf-8') as f:
            content = f.read()
            # Remove title
            lines = content.split('\n')
            for line in lines:
                if line.strip() and not line.startswith('#'):
                    add_paragraph(doc, line)
    
    doc.add_page_break()
    
    # Abstrak
    abstrak_path = os.path.join(front_matter_dir, "abstrak.md")
    if os.path.exists(abstrak_path):
        with open(abstrak_path, 'r', encoding='utf-8') as f:
            content = f.read()
            lines = content.split('\n')
            for line in lines:
                if line.strip() and not line.startswith('#'):
                    add_paragraph(doc, line)
    
    doc.add_page_break()
    
    # Daftar Isi
    daftar_isi_path = os.path.join(front_matter_dir, "daftar_isi.md")
    if os.path.exists(daftar_isi_path):
        with open(daftar_isi_path, 'r', encoding='utf-8') as f:
            content = f.read()
            add_heading(doc, "DAFTAR ISI", level=1)
            lines = content.split('\n')
            for line in lines:
                if line.strip() and not line.startswith('#'):
                    add_paragraph(doc, line)
    
    doc.add_page_break()
    
    # BAB I
    bab1_path = os.path.join(REPORT_DIR, "bab1", "pendahuluan.md")
    if os.path.exists(bab1_path):
        with open(bab1_path, 'r', encoding='utf-8') as f:
            content = f.read()
            lines = content.split('\n')
            for line in lines:
                if line.startswith('#'):
                    level = line.count('#')
                    text = line.lstrip('#').strip()
                    add_heading(doc, text, level=min(level, 3))
                elif line.strip():
                    add_paragraph(doc, line)
    
    # BAB II
    bab2_path = os.path.join(REPORT_DIR, "bab2", "kajian_pustaka.md")
    if os.path.exists(bab2_path):
        with open(bab2_path, 'r', encoding='utf-8') as f:
            content = f.read()
            lines = content.split('\n')
            for line in lines:
                if line.startswith('#'):
                    level = line.count('#')
                    text = line.lstrip('#').strip()
                    add_heading(doc, text, level=min(level, 3))
                elif line.strip():
                    add_paragraph(doc, line)
    
    # BAB III
    bab3_path = os.path.join(REPORT_DIR, "bab3", "rancangan_dan_metode.md")
    if os.path.exists(bab3_path):
        with open(bab3_path, 'r', encoding='utf-8') as f:
            content = f.read()
            lines = content.split('\n')
            for line in lines:
                if line.startswith('#'):
                    level = line.count('#')
                    text = line.lstrip('#').strip()
                    add_heading(doc, text, level=min(level, 3))
                elif line.strip():
                    add_paragraph(doc, line)
    
    # BAB IV
    bab4_path = os.path.join(REPORT_DIR, "bab4", "hasil_dan_pembahasan.md")
    if os.path.exists(bab4_path):
        with open(bab4_path, 'r', encoding='utf-8') as f:
            content = f.read()
            lines = content.split('\n')
            for line in lines:
                if line.startswith('#'):
                    level = line.count('#')
                    text = line.lstrip('#').strip()
                    add_heading(doc, text, level=min(level, 3))
                elif line.strip():
                    add_paragraph(doc, line)
    
    # BAB V
    bab5_path = os.path.join(REPORT_DIR, "bab5", "penutup.md")
    if os.path.exists(bab5_path):
        with open(bab5_path, 'r', encoding='utf-8') as f:
            content = f.read()
            lines = content.split('\n')
            for line in lines:
                if line.startswith('#'):
                    level = line.count('#')
                    text = line.lstrip('#').strip()
                    add_heading(doc, text, level=min(level, 3))
                elif line.strip():
                    add_paragraph(doc, line)
    
    doc.add_page_break()
    
    # Daftar Pustaka
    daftar_pustaka_path = os.path.join(REPORT_DIR, "daftar_pustaka.md")
    if os.path.exists(daftar_pustaka_path):
        with open(daftar_pustaka_path, 'r', encoding='utf-8') as f:
            content = f.read()
            add_heading(doc, "DAFTAR PUSTAKA", level=1)
            lines = content.split('\n')
            for line in lines:
                if line.strip() and not line.startswith('#'):
                    add_paragraph(doc, line, size=10)
    
    doc.add_page_break()
    
    # Lampiran
    lampiran_path = os.path.join(REPORT_DIR, "lampiran", "lampiran.md")
    if os.path.exists(lampiran_path):
        with open(lampiran_path, 'r', encoding='utf-8') as f:
            content = f.read()
            add_heading(doc, "LAMPIRAN", level=1)
            lines = content.split('\n')
            for line in lines:
                if line.startswith('#'):
                    level = line.count('#')
                    text = line.lstrip('#').strip()
                    add_heading(doc, text, level=min(level, 3))
                elif line.strip():
                    add_paragraph(doc, line)
    
    # Save document
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
